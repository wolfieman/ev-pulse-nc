"""Build the SEINFORMS 2026 conference .docx from the de-identified manuscript.

Run from the repo root with the project's uv environment:

    uv run python src/paper/build_seinforms_docx.py

Reads the (gitignored, double-blind-scrubbed) working manuscript at
``conferences/seinforms-2026/seinforms_manuscript.md`` and produces a
SEINFORMS-formatted ``seinforms_manuscript.docx``. The .docx is the editable
intermediate; the manual PDF export is the canonical submission artifact
(handled later, inter-agency task ``seinforms-005``). Output defaults to the
shared ``inter-agency/seinforms/output/`` directory when present, else to the
conferences folder beside the source.

This is the *formatting* tool (inter-agency task ``seinforms-003``). It applies
SEINFORMS typography only. Two adjacent concerns are deliberately out of scope:

  - Author-year -> bracketed-numeric ``[n]`` citation conversion and the
    alphabetical numbered REFERENCES list (task ``seinforms-004``, codex).
    SEINFORMS accepts author-year citations, so this script already produces a
    submittable DOCX; once the source is converted to ``[n]`` form it renders
    that faithfully with no script change.
  - Content condensation, the <=200-word in-paper abstract swap, and moving
    Appendix A before References (CFP: appendices precede references). Those are
    source-text edits made upstream of this script.

SEINFORMS format applied (per conferences/seinforms-2026/cfp-requirements.md):
  - 11 pt Times New Roman, single column, 1" margins all sides.
  - Body single-spaced and justified; flush-left paragraphs, no first-line
    indent; one blank line between paragraphs and before/after headings.
  - Title: bold, ALL CAPS, centered at the top of page 1; one blank line after.
  - Level-1 headings (manuscript ``##``): bold, centered, ALL CAPS.
  - Level-2 headings (manuscript ``###``): bold, flush left, mixed case.
  - Level-3 headings (manuscript ``####``): bold italic, flush left, mixed case.
  - Pipe tables -> Word tables; the seven high-impact figures embedded near
    their first citation.
  - References rendered with a hanging indent (works for both author-year and
    numbered entries).
  - Double-blind: no author, affiliation, advisor, or date is emitted; the
    manuscript source is already scrubbed (task ``seinforms-002``).

Copyright © 2026 Wolfgang Sanyer
Licensed under the Polyform Noncommercial License 1.0.0 (see LICENSE).
"""

from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from evpulse.paths import PROJECT_ROOT as REPO_ROOT

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SRC = REPO_ROOT / "conferences" / "seinforms-2026" / "seinforms_manuscript.md"
FIGURES_DIR = REPO_ROOT / "output" / "figures"

# Prefer the shared inter-agency output dir (per the seinforms-003 handoff);
# fall back to the conferences folder beside the source if it is absent.
_SHARED_WS = REPO_ROOT.parent / "inter-agency" / "seinforms"
_OUT_DIR = (_SHARED_WS / "output") if _SHARED_WS.is_dir() else SRC.parent
OUT = _OUT_DIR / "seinforms_manuscript.docx"

# Owner-approved: keep the original capstone title (known, accepted double-blind
# exposure to preserve project branding). Rendered ALL CAPS for the SEINFORMS L1
# title style.
TITLE = "EV Pulse NC: Data-Driven EV Infrastructure Gap Analysis for North Carolina"

# ---------------------------------------------------------------------------
# Typography constants
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
MONO_FONT = "Consolas"
FONT_SIZE = Pt(11)
PARA_SPACE = Pt(11)  # one blank 11 pt line between blocks ("skip one line")
BULLET_SPACE = Pt(4)  # tighter spacing within list runs
REF_SPACE = Pt(6)  # reference entries: compact but separated

# ---------------------------------------------------------------------------
# Figure embedding: figure number -> (PNG filename, caption title).
# Each figure is embedded immediately after the FIRST paragraph that cites it
# (matched via \bFigure N\b). The manuscript cites exactly these seven.
# ---------------------------------------------------------------------------
FIGURE_FILES = {
    24: "fig-24-heatmap-mecklenburg.png",
    33: "fig-33-theil-decomposition.png",
    36: "fig-36-demand-comparison.png",
    42: "fig-42-stations-justice40-overlay.png",
    43: "fig-43-nevi-priority-scores.png",
    44: "fig-44-validation-scatter.png",
    45: "fig-45-equity-utilization-archetypes.png",
}

FIGURE_TITLES = {
    24: "Mecklenburg County: Charging-Station Density Heat Map",
    33: (
        "Theil-T Decomposition of EV Charging Infrastructure Inequality: "
        "Within-County vs Between-County"
    ),
    36: "Workplace Charging Demand: LODES-Adjusted vs Unadjusted Estimates",
    42: "Charging Stations Overlaid on Justice40-Designated Disadvantaged Tracts",
    43: "NEVI Priority Scores: Top-10 NC Counties",
    44: (
        "Forecast Validation: Actual Versus Predicted BEV Registrations "
        "Across 400 County-Month Observations"
    ),
    45: "Equity-Utilization Quadrant: Three County Archetypes",
}

FIGURE_REF_RE = re.compile(r"\bFigure (\d+)\b")
FIGURE_WIDTH = Inches(6.0)


# ---------------------------------------------------------------------------
# Run helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, name: str = FONT_NAME) -> None:
    run.font.name = name
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def _add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    mono: bool = False,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    _set_run_font(run, MONO_FONT if mono else FONT_NAME)


# Inline tokenizer: **bold**, *italic*, `code`, non-greedy. Bold before italic.
_INLINE_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)"
    r"|(\*(?P<italic>[^*\n]+?)\*)"
    r"|(`(?P<code>[^`\n]+?)`)"
)


def add_inline_runs(paragraph, text: str, *, base_italic: bool = False) -> None:
    """Add runs to `paragraph` interpreting **bold**, *italic*, and `code`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos : m.start()], italic=base_italic)
        if m.group("bold") is not None:
            _add_run(paragraph, m.group("bold"), bold=True, italic=base_italic)
        elif m.group("italic") is not None:
            _add_run(paragraph, m.group("italic"), italic=True)
        elif m.group("code") is not None:
            _add_run(paragraph, m.group("code"), mono=True, italic=base_italic)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], italic=base_italic)


def _new_paragraph(doc: Document, *, space_after=PARA_SPACE):
    """A single-spaced paragraph with no first-line indent and a trailing gap."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.first_line_indent = Inches(0)
    return p


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------
def configure_document(doc: Document) -> None:
    """1" margins; Times New Roman 11 pt; single-spaced Normal with line gap."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = PARA_SPACE


def clear_metadata(doc: Document) -> None:
    """Blank the document core properties for double-blind review.

    python-docx seeds ``author`` to ``"python-docx"``; clearing author, title,
    and the edit-history fields keeps identifying metadata out of the DOCX (and,
    by extension, the PDF export). Final metadata scrubbing still happens at
    export time (task ``seinforms-005``).
    """
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    cp.category = ""


def add_page_numbers(doc: Document) -> None:
    """Top-right page numbers on every page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run()
    _set_run_font(run)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


# ---------------------------------------------------------------------------
# Title and heading builders
# ---------------------------------------------------------------------------
def add_title(doc: Document) -> None:
    """SEINFORMS title: bold, ALL CAPS, centered, top of page 1."""
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, TITLE.upper(), bold=True)


def add_l1(doc: Document, text: str) -> None:
    """Level 1: bold, centered, ALL CAPS."""
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text.upper(), bold=True)


def add_l2(doc: Document, text: str) -> None:
    """Level 2: bold, flush left, mixed case."""
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, text, bold=True)


def add_l3(doc: Document, text: str) -> None:
    """Level 3: bold italic, flush left, mixed case."""
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, text, bold=True, italic=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Body: single-spaced, justified, flush left, no indent."""
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = BULLET_SPACE
    add_inline_runs(p, text)


def add_reference_paragraph(doc: Document, text: str) -> None:
    """Reference entry: hanging indent, single-spaced, justified."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = REF_SPACE
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)  # hanging indent
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text)


def add_figure(doc: Document, fig_num: int) -> bool:
    """Embed a figure near its citation. Returns True if embedded."""
    png_name = FIGURE_FILES.get(fig_num)
    title_text = FIGURE_TITLES.get(fig_num)
    if not png_name or not title_text:
        print(f"WARNING: no figure mapping for Figure {fig_num}, skipping")
        return False

    png_path = FIGURES_DIR / png_name
    if not png_path.is_file():
        print(f"WARNING: figure {png_path} not found, skipping")
        return False

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    img_para.paragraph_format.first_line_indent = Inches(0)
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(2)
    img_para.add_run().add_picture(str(png_path), width=FIGURE_WIDTH)

    # Caption: "Figure N." (bold) + title (regular), centered, single-spaced.
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = PARA_SPACE
    _add_run(cap, f"Figure {fig_num}. ", bold=True)
    _add_run(cap, title_text)
    return True


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------
def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a 2D list of cell strings as a Word table with light grid style."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if r_idx == 0:
                stripped = cell_text.strip()
                if stripped.startswith("**") and stripped.endswith("**"):
                    stripped = stripped[2:-2]
                _add_run(p, stripped, bold=True)
            else:
                add_inline_runs(p, cell_text.strip())

    # "Skip one line" after the table.
    _new_paragraph(doc, space_after=Pt(0))


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
def parse_markdown(md_text: str) -> list[dict]:
    """Convert manuscript markdown into a list of block dicts.

    Everything up to and including the first `---` (the title-page header: the
    font-note comment, the `#` title, the redacted author line, and the keywords
    line) is dropped; the title is composed separately by `add_title`. The first
    parsed block is therefore `## Abstract`.

    Block types: h1/h2/h3/h4, paragraph, bullet, table, hr.
    """
    all_lines = md_text.splitlines()
    first_hr = next(
        (idx for idx, ln in enumerate(all_lines) if ln.strip() == "---"),
        None,
    )
    lines = all_lines[first_hr + 1 :] if first_hr is not None else all_lines
    blocks: list[dict] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        if stripped.startswith("#### "):
            blocks.append({"type": "h4", "text": stripped[5:].strip()})
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:].strip()})
            i += 1
            continue

        # Pipe table: header row + separator row + body rows
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < n:
            sep = lines[i + 1].strip()
            sep_chars = set(sep.replace("|", "").strip())
            if sep.startswith("|") and sep_chars <= set("-: "):
                table_rows: list[list[str]] = [_split_table_row(stripped)]
                i += 2
                while i < n:
                    row_line = lines[i].strip()
                    if row_line.startswith("|") and row_line.endswith("|"):
                        table_rows.append(_split_table_row(row_line))
                        i += 1
                    else:
                        break
                blocks.append({"type": "table", "rows": table_rows})
                continue

        if stripped.startswith("- "):
            blocks.append({"type": "bullet", "text": stripped[2:].strip()})
            i += 1
            continue

        # Paragraph: collect contiguous non-blank, non-special lines.
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt_stripped = lines[i].rstrip().strip()
            if not nxt_stripped:
                break
            if (
                nxt_stripped.startswith("#")
                or nxt_stripped == "---"
                or nxt_stripped.startswith("- ")
                or (nxt_stripped.startswith("|") and nxt_stripped.endswith("|"))
            ):
                break
            para_lines.append(nxt_stripped)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def _split_table_row(line: str) -> list[str]:
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------
def _is_references_heading(text: str) -> bool:
    return "REFERENCES" in text.upper()


def render_blocks(doc: Document, blocks: list[dict]) -> dict:
    """Render parsed blocks into the document. Returns a stats dict."""
    stats = {
        "tables": 0,
        "bullets": 0,
        "l1": 0,
        "l2": 0,
        "l3": 0,
        "paragraphs": 0,
        "figures": 0,
    }

    in_references = False
    pending_figures: set[int] = set(FIGURE_FILES.keys())

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            # Section separators -> continuous flow (no page breaks).
            continue

        if btype in ("h1", "h2"):
            # `#` survivors (e.g. "# Appendices") and all `##` are SEINFORMS L1.
            text = block["text"]
            in_references = _is_references_heading(text)
            add_l1(doc, text)
            stats["l1"] += 1

        elif btype == "h3":
            add_l2(doc, block["text"])
            stats["l2"] += 1

        elif btype == "h4":
            add_l3(doc, block["text"])
            stats["l3"] += 1

        elif btype == "table":
            add_markdown_table(doc, block["rows"])
            stats["tables"] += 1

        elif btype == "bullet":
            add_bullet_paragraph(doc, block["text"])
            stats["bullets"] += 1

        elif btype == "paragraph":
            text = block["text"]
            # Under REFERENCES: entries get a hanging indent; the lone
            # "Additional Sources Consulted" explanatory paragraph stays body.
            if in_references and not text.startswith("The following works informed"):
                add_reference_paragraph(doc, text)
                stats["paragraphs"] += 1
                continue

            add_body_paragraph(doc, text)
            stats["paragraphs"] += 1

            # Embed any still-pending figures cited in this paragraph, in first
            # appearance order, then drop them from the pending set.
            if pending_figures:
                seen_here: list[int] = []
                for m in FIGURE_REF_RE.finditer(text):
                    fig_num = int(m.group(1))
                    if fig_num in pending_figures and fig_num not in seen_here:
                        seen_here.append(fig_num)
                for fig_num in seen_here:
                    if add_figure(doc, fig_num):
                        stats["figures"] += 1
                    pending_figures.discard(fig_num)

    return stats


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    print(f"Read source: {SRC} ({len(md_text):,} chars)")

    blocks = parse_markdown(md_text)
    print(f"Parsed blocks: {len(blocks)}")

    doc = Document()
    configure_document(doc)
    clear_metadata(doc)
    add_page_numbers(doc)

    print("Rendering title + body (single-column, continuous flow)...")
    add_title(doc)
    stats = render_blocks(doc, blocks)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    size = OUT.stat().st_size
    print()
    print(f"Wrote: {OUT}")
    print(f"Size:  {size:,} bytes ({size / 1024:.1f} KB)")
    print()
    print("Block stats:")
    for k, v in stats.items():
        print(f"  {k:>12}: {v}")
    expected_figs = len(FIGURE_FILES)
    embedded = stats.get("figures", 0)
    print(f"  figures embedded: {embedded} / {expected_figs}")
    if embedded != expected_figs:
        print(
            f"  NOTE: {expected_figs - embedded} figure(s) not embedded "
            f"(citation not found in source or PNG missing)."
        )


if __name__ == "__main__":
    main()
