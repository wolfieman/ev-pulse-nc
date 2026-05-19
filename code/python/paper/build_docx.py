"""Build the APA 7 student-paper .docx for the EV Pulse NC capstone manuscript.

Run from the repo root with the project's uv environment:

    uv run python code/python/paper/build_docx.py

Reads paper/manuscript.md and produces paper/manuscript.docx. The .docx is
gitignored (`*.docx`); only this script is intended to live under version
control.

Manuscript features handled:
  - Four heading levels (#, ##, ###, ####) mapped to APA title / L1 / L2 / L3
  - Pipe-style markdown tables rendered as Word tables (Light Grid Accent 1)
  - Inline **bold**, *italic*, and `code` runs
  - Unicode math / Greek / en-dash characters preserved as-is
  - `-` bullet lists rendered as Word bulleted paragraphs
  - URLs rendered as plain text (no hyperlink conversion)
  - References section with hanging 0.5" indent, italic journal/book titles
  - Appendices on new pages
  - APA 7 student-paper title page (centered + bold title, byline, advisor, date)
  - Abstract and Keywords on the page following the title page
  - Page numbers in the top-right header on every page
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
# Script lives at <repo>/code/python/paper/build_docx.py; walk up 3 parents
# to reach the repo root, then point at the paper/ and output/ siblings.
REPO_ROOT = Path(__file__).resolve().parents[3]
SRC = REPO_ROOT / "paper" / "manuscript.md"
OUT = REPO_ROOT / "paper" / "manuscript.docx"
FIGURES_DIR = REPO_ROOT / "output" / "figures"

FONT_NAME = "Times New Roman"
MONO_FONT = "Consolas"
FONT_SIZE = Pt(12)

# ---------------------------------------------------------------------------
# Figure embedding configuration
# ---------------------------------------------------------------------------
# Map figure number -> (PNG filename, italic caption text).
# Each figure is embedded immediately after the FIRST paragraph that cites it
# (matched via the regex \bFigure N\b). Subsequent references to the same
# figure number do not trigger another embed.
FIGURE_TITLES = {
    24: "Mecklenburg County: Charging-Station Density Heat Map",
    33: "Theil-T Decomposition of EV Charging Infrastructure Inequality: Within-County vs Between-County",
    36: "Workplace Charging Demand: LODES-Adjusted vs Unadjusted Estimates",
    42: "Charging Stations Overlaid on Justice40-Designated Disadvantaged Tracts",
    43: "NEVI Priority Scores: Top-10 NC Counties",
    44: "Forecast Validation: Actual Versus Predicted BEV Registrations Across 400 County-Month Observations",
    45: "Equity-Utilization Quadrant: Three County Archetypes",
}

FIGURE_FILES = {
    24: "fig-24-heatmap-mecklenburg.png",
    33: "fig-33-theil-decomposition.png",
    36: "fig-36-demand-comparison.png",
    42: "fig-42-stations-justice40-overlay.png",
    43: "fig-43-nevi-priority-scores.png",
    44: "fig-44-validation-scatter.png",
    45: "fig-45-equity-utilization-archetypes.png",
}

FIGURE_REF_RE = re.compile(r"\bFigure (\d+)\b")
FIGURE_WIDTH = Inches(6.0)


# ---------------------------------------------------------------------------
# Run helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, name: str = FONT_NAME) -> None:
    run.font.name = name
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             mono: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    _set_run_font(run, MONO_FONT if mono else FONT_NAME)


# Inline tokenizer: handles **bold**, *italic*, and `code` non-greedily.
# Order matters: bold must be tried before italic (** vs *).
_INLINE_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)"
    r"|(\*(?P<italic>[^*\n]+?)\*)"
    r"|(`(?P<code>[^`\n]+?)`)"
)


def add_inline_runs(paragraph, text: str, *, base_italic: bool = False) -> None:
    """Add runs to `paragraph` interpreting **bold**, *italic*, and `code`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], italic=base_italic)
        if m.group("bold") is not None:
            _add_run(paragraph, m.group("bold"), bold=True, italic=base_italic)
        elif m.group("italic") is not None:
            _add_run(paragraph, m.group("italic"), italic=True)
        elif m.group("code") is not None:
            _add_run(paragraph, m.group("code"), mono=True, italic=base_italic)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], italic=base_italic)


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------
def configure_document(doc: Document) -> None:
    """1" margins, Times New Roman 12pt, double-spaced Normal style."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_page_numbers(doc: Document) -> None:
    """Top-right page numbers, every page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run()
    _set_run_font(run)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_paragraph(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


# ---------------------------------------------------------------------------
# Heading and body paragraph builders
# ---------------------------------------------------------------------------
def add_h1(doc: Document, text: str) -> None:
    """APA Level 1: centered, bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(p, text, bold=True)


def add_h2(doc: Document, text: str) -> None:
    """APA Level 2: flush-left, bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(p, text, bold=True)


def add_h3(doc: Document, text: str) -> None:
    """APA Level 3: flush-left, bold italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(p, text, bold=True, italic=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.first_line_indent = Inches(0.5)
    add_inline_runs(p, text)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    add_inline_runs(p, text)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)  # hanging indent
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(p, text)


def add_figure(doc: Document, fig_num: int) -> bool:
    """Embed a figure image with APA 7 caption.

    Returns True if the image was embedded, False if the PNG was missing.
    The image is added as a centered paragraph at 6.0" width (height
    auto-scales). The APA 7 caption follows in two centered paragraphs:
    (1) bold "Figure N" label, (2) italic title.
    """
    png_name = FIGURE_FILES.get(fig_num)
    title_text = FIGURE_TITLES.get(fig_num)
    if not png_name or not title_text:
        print(f"WARNING: no figure mapping for Figure {fig_num}, skipping")
        return False

    png_path = FIGURES_DIR / png_name
    if not png_path.is_file():
        print(f"WARNING: figure {png_path} not found, skipping")
        return False

    # Image paragraph: centered, no indent, single-spaced so the caption hugs
    # the image rather than floating a full double-spaced gap below it.
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    img_para.paragraph_format.first_line_indent = Inches(0)
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(0)
    run = img_para.add_run()
    run.add_picture(str(png_path), width=FIGURE_WIDTH)

    # APA 7 caption — "Figure N" label paragraph: centered, bold, no indent
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    label_para.paragraph_format.first_line_indent = Inches(0)
    label_para.paragraph_format.space_before = Pt(2)
    label_para.paragraph_format.space_after = Pt(0)
    _add_run(label_para, f"Figure {fig_num}", bold=True)

    # APA 7 caption — title paragraph: centered, italic, no indent
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    title_para.paragraph_format.first_line_indent = Inches(0)
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(6)
    _add_run(title_para, title_text, italic=True)
    return True


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------
def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a 2D list of cell strings as a Word table with light grid style."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            # Clear the default paragraph then add our own with inline formatting
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # Header row gets bold
            if r_idx == 0:
                # Strip surrounding ** if present, then bold the whole cell
                stripped = cell_text.strip()
                if stripped.startswith("**") and stripped.endswith("**"):
                    stripped = stripped[2:-2]
                _add_run(p, stripped, bold=True)
            else:
                add_inline_runs(p, cell_text.strip())


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
TITLE_PAGE_END_LINE = 12  # Inclusive 1-indexed cutoff for the title-page block


def parse_markdown(md_text: str) -> list[dict]:
    """Convert manuscript markdown into a list of block dicts.

    The title-page header (manuscript.md lines 1-12: HTML font comment, the
    `#` title, byline, affiliation, advisor, date, Keywords) is composed
    manually in `add_title_page` and is therefore stripped from the block
    stream here. The first horizontal rule (line 12) marks the end of the
    title-page region, and `## Abstract` is the first parsed block.

    Block types:
      {"type": "h1" | "h2" | "h3" | "h4", "text": str}
      {"type": "paragraph", "text": str}
      {"type": "bullet", "text": str}
      {"type": "table", "rows": list[list[str]]}
      {"type": "hr"}
      {"type": "blank"}  (used only between bullets if needed)
    """
    all_lines = md_text.splitlines()
    # Drop everything up to and including the first `---` separator: that
    # entire region is the manually-composed title-page header.
    first_hr = next(
        (idx for idx, ln in enumerate(all_lines) if ln.strip() == "---"),
        None,
    )
    lines = all_lines[first_hr + 1:] if first_hr is not None else all_lines
    blocks: list[dict] = []
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Skip HTML comments (only single-line forms present in the manuscript)
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            blocks.append({"type": "h4", "text": stripped[5:].strip()})
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:].strip()})
            i += 1
            continue

        # Pipe table: header row + separator row + body rows
        if stripped.startswith("|") and stripped.endswith("|"):
            # Lookahead: confirm next line is a separator (---|---|...)
            if i + 1 < n:
                sep = lines[i + 1].strip()
                if sep.startswith("|") and set(sep.replace("|", "").strip()) <= set("-: "):
                    # Collect rows
                    table_rows: list[list[str]] = []
                    # Header
                    table_rows.append(_split_table_row(stripped))
                    i += 2  # skip header + separator
                    while i < n:
                        row_line = lines[i].strip()
                        if row_line.startswith("|") and row_line.endswith("|"):
                            table_rows.append(_split_table_row(row_line))
                            i += 1
                        else:
                            break
                    blocks.append({"type": "table", "rows": table_rows})
                    continue

        # Bullet list item
        if stripped.startswith("- "):
            blocks.append({"type": "bullet", "text": stripped[2:].strip()})
            i += 1
            continue

        # Paragraph: collect contiguous non-blank, non-special lines
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            if (nxt_stripped.startswith("#") or nxt_stripped == "---"
                    or nxt_stripped.startswith("- ")
                    or (nxt_stripped.startswith("|") and nxt_stripped.endswith("|"))):
                break
            para_lines.append(nxt_stripped)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def _split_table_row(line: str) -> list[str]:
    # Strip leading/trailing pipes, then split on remaining pipes
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


# ---------------------------------------------------------------------------
# Title page (rendered separately from the body block stream)
# ---------------------------------------------------------------------------
def add_title_page(doc: Document) -> None:
    """APA 7 student-paper title page, manually composed from manuscript header."""
    # Push title into upper third of the page
    for _ in range(4):
        add_blank_paragraph(doc)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(
        title_para,
        "EV Pulse NC: Data-Driven EV Infrastructure Gap Analysis for North Carolina",
        bold=True,
    )

    add_blank_paragraph(doc)

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(byline, "Wolfgang Sanyer", bold=True)

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(
        affiliation,
        "Fayetteville State University · Broadwell College of Business and Economics",
    )

    advisor = doc.add_paragraph()
    advisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    advisor.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(advisor, "Faculty advisor: Dr. Majed Al-Ghandour", italic=True)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _add_run(date_p, "Date: [Stage 1 submission date]", italic=True)


# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------
def render_blocks(doc: Document, blocks: list[dict]) -> dict:
    """Render parsed blocks into the document, with structure-aware decisions.

    Returns a stats dict with counts for the build report.
    """
    stats = {
        "tables": 0, "bullets": 0, "h1": 0, "h2": 0, "h3": 0,
        "paragraphs": 0, "figures": 0,
    }

    in_references = False
    in_appendix = False
    in_abstract = False
    keywords_inserted = False
    prev_block_type: str | None = None

    # Track which figures have already been embedded so subsequent textual
    # references don't trigger a duplicate insertion.
    embedded_figures: set[int] = set()
    pending_figures: set[int] = set(FIGURE_TITLES.keys())

    KEYWORDS_PREFIX = "Keywords:"
    KEYWORDS_TEXT = (
        "**Keywords:** EV charging infrastructure, NEVI Formula Program, "
        "Justice40, Theil decomposition, infrastructure equity, North Carolina"
    )

    # We need to know when we hit "## 12. References" so subsequent paragraphs
    # use hanging-indent reference style instead of body-indent style. The
    # sub-headings "### Cited Works" and "### Additional Sources Consulted"
    # remain L2 headings; intro paragraphs to "Additional Sources Consulted"
    # render as bodies. The simplest rule: under §12, any plain paragraph
    # that begins with an author-like pattern (uppercase letter then
    # punctuation) is treated as a reference; the introductory paragraph
    # ("The following works informed...") is short and benefits from
    # hanging-indent too, so we apply hanging indent uniformly to all
    # paragraphs under §12 except the first explanatory paragraph.

    for idx, block in enumerate(blocks):
        btype = block["type"]

        if btype == "hr":
            # Treat top-level horizontal rules as section separators rather
            # than page breaks; they appear between every major section in the
            # manuscript and a page break per rule would balloon the page count.
            prev_block_type = btype
            continue

        if btype == "h1":
            # `# EV Pulse NC...` is the title (already rendered as title page).
            # `# Appendices` is a section divider before the two appendices.
            text = block["text"]
            if text.startswith("EV Pulse NC"):
                prev_block_type = btype
                continue
            if text == "Appendices":
                # Start a fresh page; the appendix H2 below will be on its own page
                add_page_break(doc)
                prev_block_type = btype
                continue
            # Other H1 (none expected) — render as L1 heading
            add_h1(doc, text)
            stats["h1"] += 1

        elif btype == "h2":
            text = block["text"]
            # New page before each top-level section except "Abstract" (which
            # follows the title page after a page break we already inserted).
            if text == "Abstract":
                in_abstract = True
                # page break already inserted before body rendering
            else:
                # Leaving the Abstract — append the Keywords paragraph before
                # the section break, so it renders at the bottom of the
                # Abstract page (APA 7 convention).
                if in_abstract and not keywords_inserted:
                    add_body_paragraph(doc, KEYWORDS_TEXT)
                    keywords_inserted = True
                    in_abstract = False
                if text.startswith("Appendix"):
                    add_page_break(doc)
                    in_appendix = True
                elif text == "12. References":
                    add_page_break(doc)
                    in_references = True
                else:
                    add_page_break(doc)

            add_h1(doc, text)
            stats["h1"] += 1

        elif btype == "h3":
            add_h2(doc, block["text"])
            stats["h2"] += 1

        elif btype == "h4":
            add_h3(doc, block["text"])
            stats["h3"] += 1

        elif btype == "table":
            add_markdown_table(doc, block["rows"])
            stats["tables"] += 1

        elif btype == "bullet":
            add_bullet_paragraph(doc, block["text"])
            stats["bullets"] += 1

        elif btype == "paragraph":
            text = block["text"]
            if in_references:
                # Under §12: treat the explanatory paragraph in "Additional
                # Sources Consulted" as a body paragraph; everything else is
                # a reference entry. The explanatory paragraph is the only
                # one that starts with "The following works informed".
                if text.startswith("The following works informed"):
                    add_body_paragraph(doc, text)
                    stats["paragraphs"] += 1
                else:
                    add_reference_paragraph(doc, text)
                    stats["paragraphs"] += 1
            else:
                add_body_paragraph(doc, text)
                stats["paragraphs"] += 1

                # After rendering the paragraph, check whether it cites any
                # of the still-pending figures. Insert each one immediately
                # after the paragraph (in the order they first appear in the
                # text), then mark them as embedded so duplicate citations
                # downstream are ignored.
                if pending_figures:
                    seen_here: list[int] = []
                    for m in FIGURE_REF_RE.finditer(text):
                        fig_num = int(m.group(1))
                        if fig_num in pending_figures and fig_num not in seen_here:
                            seen_here.append(fig_num)
                    for fig_num in seen_here:
                        if add_figure(doc, fig_num):
                            stats["figures"] += 1
                            embedded_figures.add(fig_num)
                        # Whether the image embedded or not (e.g. missing
                        # PNG), drop it from the pending set so we don't
                        # retry on every subsequent citation.
                        pending_figures.discard(fig_num)

        prev_block_type = btype

    return stats


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    print(f"Read source: {SRC} ({len(md_text):,} chars)")

    blocks = parse_markdown(md_text)
    print(f"Parsed blocks: {len(blocks)}")

    doc = Document()
    configure_document(doc)
    add_page_numbers(doc)

    print("Building title page...")
    add_title_page(doc)
    add_page_break(doc)

    print("Rendering body blocks...")
    stats = render_blocks(doc, blocks)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    size = OUT.stat().st_size
    print()
    print(f"Wrote: {OUT}")
    print(f"Size:  {size:,} bytes ({size / 1024:.1f} KB)")
    print()
    print("Block stats:")
    for k, v in stats.items():
        print(f"  {k:>12}: {v}")
    expected_figs = len(FIGURE_TITLES)
    print(f"  figures embedded: {stats.get('figures', 0)} / {expected_figs}")


if __name__ == "__main__":
    main()
